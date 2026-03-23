import io
import docx
from PyPDF2 import PdfReader
from src.utils import get_logger

logger = get_logger(__name__)

def extract_text_from_pdf(file_path_or_bytes):
    try:
        pdf_file = io.BytesIO(file_path_or_bytes) if isinstance(file_path_or_bytes, bytes) else open(file_path_or_bytes, 'rb')
        reader = PdfReader(pdf_file)
        text_chunks = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text = page_text.replace('\n\n', '<PARAGRAPH_BREAK>')
                text = text.replace('\n', ' ')
                paragraphs = [p.strip() for p in text.split('<PARAGRAPH_BREAK>') if p.strip()]
                text_chunks.extend(paragraphs)
        if not isinstance(file_path_or_bytes, bytes):
            pdf_file.close()
        return "\n\n".join(text_chunks)
    except Exception as e:
        logger.error(f"Failed to extract PDF: {e}")
        return ""

def extract_text_from_docx(file_path_or_bytes):
    try:
        doc = docx.Document(io.BytesIO(file_path_or_bytes)) if isinstance(file_path_or_bytes, bytes) else docx.Document(file_path_or_bytes)
        text_chunks = []
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    text_chunks.append(f"\n[{para.style.name.upper()}] {para.text.strip()}")
                else:
                    text_chunks.append(para.text.strip())
        return "\n".join(text_chunks)
    except Exception as e:
        logger.error(f"Failed to extract DOCX: {e}")
        return ""
